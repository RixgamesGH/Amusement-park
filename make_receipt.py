from datetime import datetime

from jinja2 import Template
from docx import Document
from docx.shared import Pt


class Receipt:

    def __init__(self):
        self.template_str = """
{%- for item, price in items.items() %}
{{ item }}:
                                                          €{{ '%0.2f' % price|round(2) }}
{% endfor %}
---------------------------------------------------
Total:                                               €{{ '%0.2f' % total|round(2) }}
        """

    def _generate_receipt_template(self, items):

        template = Template(self.template_str)

        total = sum(items.values())

        return template.render(items=items, total=total)

    def save_receipt_to_docx(self, receipt_id, items, park_name):
        doc = Document()
        receipt_text = self._generate_receipt_template(items)

        today = datetime.now()
        date = today.strftime("%B %d, %Y %H:%M")

        # Set page size
        section = doc.sections[0]
        section.page_width = Pt(252)  # Pt(72) is 1 inch
        section.page_height = Pt(288)

        # Set margins
        section.left_margin = Pt(36)
        section.right_margin = Pt(0)
        section.top_margin = Pt(14.4)
        section.bottom_margin = Pt(0)

        doc.add_heading(f"Receipt {park_name}", 2)

        # Add receipt content
        text = doc.add_paragraph()
        text.add_run(f"Date: {date}\tID: {str(receipt_id)}\n").font.size = Pt(6)
        text.add_run(receipt_text).font.size = Pt(10)
        text.paragraph_format.line_spacing = Pt(10)

        # Save the document
        filename = f"receipt{receipt_id}.docx"
        doc.save(filename)
        print(f"Receipt saved to {filename}")
